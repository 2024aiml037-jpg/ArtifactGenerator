"""
Ingestion Layer - Handles multiple data sources (PDF, Word, Confluence, DB, Code)
Converts all sources into structured JSON with metadata
"""

import os
import tempfile
import json
from typing import List, Optional
from datetime import datetime
import logging

from langchain.document_loaders import TextLoader, PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

from models.schemas import (
    IngestedDocument, Metadata, SourceType
)

logger = logging.getLogger(__name__)


class IngestionLayer:
    """Handles ingestion from multiple data sources"""

    def __init__(self):
        self.supported_formats = {'.pdf', '.txt', '.docx', '.json', '.xlsx'}
        self.document_counter = 0
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )

    def ingest_pdf(self, file_path: str, metadata: Optional[Metadata] = None) -> IngestedDocument:
        """Ingest PDF file"""
        try:
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            content = "\n".join([doc.page_content for doc in documents])
            
            if metadata is None:
                metadata = Metadata(
                    source=SourceType.PDF,
                    filename=os.path.basename(file_path),
                    timestamp=datetime.utcnow()
                )
            
            doc_id = f"doc_{self.document_counter}_{int(datetime.utcnow().timestamp())}"
            self.document_counter += 1
            
            return IngestedDocument(
                document_id=doc_id,
                content=content,
                metadata=metadata,
                chunks=[doc.page_content for doc in documents]
            )
        except Exception as e:
            logger.error(f"Error ingesting PDF {file_path}: {str(e)}")
            raise

    def ingest_text(self, file_path: str, metadata: Optional[Metadata] = None) -> IngestedDocument:
        """Ingest plain text file"""
        try:
            loader = TextLoader(file_path)
            documents = loader.load()
            content = "".join([doc.page_content for doc in documents])
            
            if metadata is None:
                metadata = Metadata(
                    source=SourceType.TEXT,
                    filename=os.path.basename(file_path),
                    timestamp=datetime.utcnow()
                )
            
            doc_id = f"doc_{self.document_counter}_{int(datetime.utcnow().timestamp())}"
            self.document_counter += 1
            
            return IngestedDocument(
                document_id=doc_id,
                content=content,
                metadata=metadata,
                chunks=[content]
            )
        except Exception as e:
            logger.error(f"Error ingesting text file {file_path}: {str(e)}")
            raise

    def ingest_docx(self, file_path: str, metadata: Optional[Metadata] = None) -> IngestedDocument:
        """Ingest Word document"""
        if DocxDocument is None:
            raise ImportError("python-docx is not installed. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            content = "\n".join([para.text for para in doc.paragraphs])
            
            if metadata is None:
                metadata = Metadata(
                    source=SourceType.WORD,
                    filename=os.path.basename(file_path),
                    timestamp=datetime.utcnow()
                )
            
            doc_id = f"doc_{self.document_counter}_{int(datetime.utcnow().timestamp())}"
            self.document_counter += 1
            
            return IngestedDocument(
                document_id=doc_id,
                content=content,
                metadata=metadata,
                chunks=[content]
            )
        except Exception as e:
            logger.error(f"Error ingesting DOCX {file_path}: {str(e)}")
            raise

    def ingest_json(self, file_path: str, metadata: Optional[Metadata] = None) -> IngestedDocument:
        """Ingest JSON file into structured chunks by section"""
        try:
            with open(file_path, 'r') as f:
                data = json.load(f)
            
            content = json.dumps(data, indent=2)
            
            # Create meaningful chunks by top-level keys
            chunks = []
            if isinstance(data, dict):
                for key, value in data.items():
                    chunk = f"[{key}]: {json.dumps(value, indent=2)}"
                    chunks.append(chunk)
            elif isinstance(data, list):
                for i, item in enumerate(data):
                    chunk = f"[item_{i}]: {json.dumps(item, indent=2)}"
                    chunks.append(chunk)
            
            # If any chunk is too large, split further
            final_chunks = []
            for chunk in chunks:
                if len(chunk) > 1000:
                    final_chunks.extend(self.text_splitter.split_text(chunk))
                else:
                    final_chunks.append(chunk)
            
            if not final_chunks:
                final_chunks = [content]
            
            if metadata is None:
                metadata = Metadata(
                    source=SourceType.API,
                    filename=os.path.basename(file_path),
                    timestamp=datetime.utcnow()
                )
            
            doc_id = f"doc_{self.document_counter}_{int(datetime.utcnow().timestamp())}"
            self.document_counter += 1
            
            return IngestedDocument(
                document_id=doc_id,
                content=content,
                metadata=metadata,
                chunks=final_chunks
            )
        except Exception as e:
            logger.error(f"Error ingesting JSON {file_path}: {str(e)}")
            raise

    def ingest_xlsx(self, file_path: str, metadata: Optional[Metadata] = None) -> IngestedDocument:
        """Ingest Excel file into structured chunks per sheet/row-group"""
        if load_workbook is None:
            raise ImportError("openpyxl is not installed. Install with: pip install openpyxl")
        
        try:
            workbook = load_workbook(file_path)
            content_parts = []
            chunks = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                rows = list(sheet.iter_rows(values_only=True))
                
                if not rows:
                    continue
                
                # First row as header
                headers = [str(cell) if cell is not None else "" for cell in rows[0]]
                sheet_header = f"=== Sheet: {sheet_name} ==="
                header_line = " | ".join(headers)
                content_parts.append(sheet_header)
                content_parts.append(header_line)
                
                # Each data row becomes a chunk with headers as context
                for row in rows[1:]:
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    # Create a structured chunk: header-value pairs
                    row_pairs = [f"{h}: {v}" for h, v in zip(headers, row_values) if v]
                    if row_pairs:
                        chunk = f"[{sheet_name}] " + " | ".join(row_pairs)
                        chunks.append(chunk)
                    
                    row_content = " | ".join(row_values)
                    content_parts.append(row_content)
            
            content = "\n".join(content_parts)
            
            if not chunks:
                chunks = self.text_splitter.split_text(content) if content else [content]
            
            if metadata is None:
                metadata = Metadata(
                    source=SourceType.DATABASE,
                    filename=os.path.basename(file_path),
                    timestamp=datetime.utcnow()
                )
            
            doc_id = f"doc_{self.document_counter}_{int(datetime.utcnow().timestamp())}"
            self.document_counter += 1
            
            return IngestedDocument(
                document_id=doc_id,
                content=content,
                metadata=metadata,
                chunks=chunks
            )
        except Exception as e:
            logger.error(f"Error ingesting XLSX {file_path}: {str(e)}")
            raise

    def ingest_file(self, file_path: str, metadata: Optional[Metadata] = None) -> IngestedDocument:
        """Ingest file based on extension"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.ingest_pdf(file_path, metadata)
        elif file_ext == '.txt':
            return self.ingest_text(file_path, metadata)
        elif file_ext == '.docx':
            return self.ingest_docx(file_path, metadata)
        elif file_ext == '.json':
            return self.ingest_json(file_path, metadata)
        elif file_ext == '.xlsx':
            return self.ingest_xlsx(file_path, metadata)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}. Supported: {self.supported_formats}")

    def ingest_from_file_object(self, file_obj, filename: str, metadata: Optional[Metadata] = None) -> IngestedDocument:
        """Ingest from file object (via Flask upload)"""
        temp_dir = tempfile.mkdtemp()
        temp_path = os.path.join(temp_dir, filename)
        
        try:
            file_obj.save(temp_path)
            
            if metadata is None:
                file_ext = os.path.splitext(filename)[1].lower()
                source_map = {
                    '.pdf': SourceType.PDF,
                    '.docx': SourceType.WORD,
                    '.json': SourceType.API,
                    '.xlsx': SourceType.DATABASE,
                    '.txt': SourceType.TEXT,
                }
                source_type = source_map.get(file_ext, SourceType.TEXT)
                metadata = Metadata(
                    source=source_type,
                    filename=filename,
                    timestamp=datetime.utcnow()
                )
            
            return self.ingest_file(temp_path, metadata)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
            if os.path.exists(temp_dir):
                os.rmdir(temp_dir)

    def ingest_text_content(self, content: str, source_name: str, metadata: Optional[Metadata] = None) -> IngestedDocument:
        """Ingest raw text content"""
        if metadata is None:
            metadata = Metadata(
                source=SourceType.TEXT,
                filename=source_name,
                timestamp=datetime.utcnow()
            )
        
        doc_id = f"doc_{self.document_counter}_{int(datetime.utcnow().timestamp())}"
        self.document_counter += 1
        
        return IngestedDocument(
            document_id=doc_id,
            content=content,
            metadata=metadata,
            chunks=[content]
        )

    def ingest_db_schema(self, db_url: str, metadata: Optional[Metadata] = None) -> IngestedDocument:
        """Ingest database schema"""
        try:
            import sqlalchemy as sa
            
            engine = sa.create_engine(db_url)
            inspector = sa.inspect(engine)
            
            schema_info = {}
            for table_name in inspector.get_table_names():
                columns = inspector.get_columns(table_name)
                schema_info[table_name] = [
                    {
                        'name': col['name'],
                        'type': str(col['type']),
                        'nullable': col['nullable']
                    }
                    for col in columns
                ]
            
            content = json.dumps(schema_info, indent=2)
            
            if metadata is None:
                metadata = Metadata(
                    source=SourceType.DATABASE,
                    filename="database_schema",
                    timestamp=datetime.utcnow()
                )
            
            doc_id = f"doc_{self.document_counter}_{int(datetime.utcnow().timestamp())}"
            self.document_counter += 1
            
            return IngestedDocument(
                document_id=doc_id,
                content=content,
                metadata=metadata,
                chunks=[content]
            )
        except ImportError:
            raise ImportError("sqlalchemy is not installed. Install with: pip install sqlalchemy")
        except Exception as e:
            logger.error(f"Error ingesting database schema: {str(e)}")
            raise

    def ingest_code_file(self, file_path: str, metadata: Optional[Metadata] = None) -> IngestedDocument:
        """Ingest source code file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if metadata is None:
                metadata = Metadata(
                    source=SourceType.CODE,
                    filename=os.path.basename(file_path),
                    timestamp=datetime.utcnow()
                )
            
            doc_id = f"doc_{self.document_counter}_{int(datetime.utcnow().timestamp())}"
            self.document_counter += 1
            
            return IngestedDocument(
                document_id=doc_id,
                content=content,
                metadata=metadata,
                chunks=[content]
            )
        except Exception as e:
            logger.error(f"Error ingesting code file {file_path}: {str(e)}")
            raise
